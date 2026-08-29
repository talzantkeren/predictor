from __future__ import annotations

import argparse
import hashlib
import re
import sys
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "docs" / "project-book-source.md"
DEFAULT_OUTPUT = ROOT / "docs" / "project-book.docx"
FIXED_TIME = datetime(2026, 8, 27, 12, 0, tzinfo=timezone.utc)
ZIP_TIME = (2026, 8, 27, 12, 0, 0)
LINK_PATTERN = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
INLINE_PATTERN = re.compile(r"\[([^\]]+)\]\(([^)]+)\)|`([^`]+)`")
TABLE_SEPARATOR_PATTERN = re.compile(r"^:?-{3,}:?$")


@dataclass(frozen=True)
class Block:
    kind: str
    value: object


def parse_markdown(source: str) -> list[Block]:
    lines = source.splitlines()
    blocks: list[Block] = []
    paragraph: list[str] = []
    index = 0

    def flush_paragraph() -> None:
        if paragraph:
            blocks.append(Block("paragraph", " ".join(paragraph).strip()))
            paragraph.clear()

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()

        if line == "<!-- pagebreak -->":
            flush_paragraph()
            blocks.append(Block("pagebreak", None))
            index += 1
            continue
        if line.startswith("<!--"):
            flush_paragraph()
            while index < len(lines) and "-->" not in lines[index]:
                index += 1
            index += 1
            continue
        if not line:
            flush_paragraph()
            index += 1
            continue
        if line.startswith("|"):
            flush_paragraph()
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = [
                [cell.strip() for cell in candidate.strip("|").split("|")]
                for candidate in table_lines
            ]
            if len(rows) < 2 or not all(
                TABLE_SEPARATOR_PATTERN.fullmatch(cell.replace(" ", ""))
                for cell in rows[1]
            ):
                raise ValueError("Markdown table is missing a separator row")
            blocks.append(Block("table", [rows[0], *rows[2:]]))
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            blocks.append(Block(f"heading{len(heading.group(1))}", heading.group(2)))
            index += 1
            continue
        if line.startswith("- "):
            flush_paragraph()
            blocks.append(Block("bullet", line[2:].strip()))
            index += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            flush_paragraph()
            blocks.append(Block("number", re.sub(r"^\d+\.\s+", "", line)))
            index += 1
            continue
        if line.startswith("> "):
            flush_paragraph()
            blocks.append(Block("callout", line[2:].strip()))
            index += 1
            continue

        paragraph.append(line)
        index += 1

    flush_paragraph()
    return blocks


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
    font_name: str = "Arial",
    rtl: bool | None = None,
) -> None:
    run.font.name = font_name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), font_name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), font_name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), font_name)
    run_properties = run._element.get_or_add_rPr()
    language = run_properties.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        run_properties.append(language)
    language.set(qn("w:val"), "he-IL")
    if rtl is not None:
        direction = run_properties.find(qn("w:rtl"))
        if direction is None:
            direction = OxmlElement("w:rtl")
            run_properties.append(direction)
        direction.set(qn("w:val"), "1" if rtl else "0")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_hyperlink(paragraph, label: str, target: str) -> None:
    relationship_id = paragraph.part.relate_to(
        target,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1769AA")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Arial")
    run_properties.extend([fonts, color, underline])
    direction = OxmlElement("w:rtl")
    direction.set(qn("w:val"), "1" if re.search(r"[\u0590-\u05ff]", label) else "0")
    run_properties.append(direction)
    text = OxmlElement("w:t")
    text.text = label
    run.extend([run_properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_directional_runs(
    paragraph,
    text: str,
    *,
    size: float | None,
    bold: bool | None,
    color: RGBColor | None,
    font_name: str = "Arial",
) -> None:
    for chunk in re.split(r"([A-Za-z0-9][A-Za-z0-9+./:_-]*)", text):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        rtl = bool(re.search(r"[\u0590-\u05ff]", chunk))
        set_run_font(
            run,
            size=size,
            bold=bold,
            color=color,
            font_name=font_name,
            rtl=rtl,
        )


def add_rich_text(paragraph, text: str, *, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    cursor = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > cursor:
            add_directional_runs(
                paragraph,
                text[cursor : match.start()],
                size=size,
                bold=bold,
                color=color,
            )
        if match.group(1) is not None and match.group(2) is not None:
            add_hyperlink(paragraph, match.group(1), match.group(2))
        else:
            add_directional_runs(
                paragraph,
                match.group(3),
                size=size,
                bold=bold,
                color=RGBColor(0x1F, 0x4E, 0x79),
                font_name="Consolas",
            )
        cursor = match.end()
    if cursor < len(text):
        add_directional_runs(
            paragraph,
            text[cursor:],
            size=size,
            bold=bold,
            color=color,
        )


def add_page_number(paragraph) -> None:
    run = paragraph.add_run("עמוד ")
    set_run_font(run, size=9, color=RGBColor(0x5B, 0x64, 0x70))
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, display, end])


def configure_styles(document: Document) -> None:
    palette = {
        "navy": RGBColor(0x1F, 0x4E, 0x79),
        "blue": RGBColor(0x2E, 0x75, 0xB6),
        "text": RGBColor(0x20, 0x2A, 0x35),
        "muted": RGBColor(0x5B, 0x64, 0x70),
    }
    style_tokens = {
        "Normal": (11, False, palette["text"], 0, 6),
        "Title": (30, True, palette["navy"], 0, 8),
        "Subtitle": (15, False, palette["muted"], 0, 14),
        "Heading 1": (20, True, palette["blue"], 12, 5),
        "Heading 2": (15, True, palette["navy"], 9, 4),
        "Heading 3": (12, True, palette["navy"], 7, 3),
        "List Bullet": (10.5, False, palette["text"], 0, 3),
        "List Number": (10.5, False, palette["text"], 0, 3),
    }
    for name, (size, bold, color, before, after) in style_tokens.items():
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.12
        style.paragraph_format.keep_with_next = name.startswith("Heading")
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        for attribute in ("ascii", "hAnsi", "cs"):
            r_fonts.set(qn(f"w:{attribute}"), "Arial")


def configure_section(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    header = section.header
    header_paragraph = header.paragraphs[0]
    set_paragraph_rtl(header_paragraph)
    add_rich_text(
        header_paragraph,
        "Predictor1 | ספר פרויקט",
        size=9,
        bold=True,
        color=RGBColor(0x5B, 0x64, 0x70),
    )

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_rich_text(
        footer_paragraph,
        "Predictor1 | RUNI 2026 | v1.4 | 30.8.2026 | page ",
        size=9,
        color=RGBColor(0x5B, 0x64, 0x70),
    )
    run = footer_paragraph.add_run()
    set_run_font(run, size=9, color=RGBColor(0x5B, 0x64, 0x70), rtl=False)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, display, end])


def create_numbering_instance(document: Document) -> int:
    numbering = document.part.numbering_part.element
    style_num_id = document.styles["List Number"].element.pPr.numPr.numId.val
    source_number = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == style_num_id
    )
    abstract_num_id = source_number.find(qn("w:abstractNumId")).get(qn("w:val"))
    next_num_id = max(int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))) + 1
    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    number.extend([abstract, override])
    numbering.append(number)
    return next_num_id


def apply_numbering(paragraph, num_id: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    num_properties = paragraph_properties.get_or_add_numPr()
    level = num_properties.get_or_add_ilvl()
    level.val = 0
    number = num_properties.get_or_add_numId()
    number.val = num_id


def set_table_geometry(table, column_widths: list[int]) -> None:
    table.autofit = False
    table_element = table._tbl
    table_properties = table_element.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(sum(column_widths)))
    table_width.set(qn("w:type"), "dxa")
    table_indent = OxmlElement("w:tblInd")
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    table_properties.append(table_indent)
    if table_properties.find(qn("w:bidiVisual")) is None:
        table_properties.append(OxmlElement("w:bidiVisual"))

    grid = table_element.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in column_widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = column_widths[index]
            cell.width = Inches(width / 1440)
            tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, rows: list[list[str]]) -> None:
    column_count = len(rows[0])
    if any(len(row) != column_count for row in rows):
        raise ValueError("Markdown table has inconsistent column counts")
    if column_count == 3:
        widths = [1900, 3300, 3920]
    elif column_count == 2:
        widths = [2700, 6420]
    else:
        widths = [9120 // column_count] * column_count
        widths[-1] += 9120 - sum(widths)

    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header_row_properties = table.rows[0]._tr.get_or_add_trPr()
    header_row_properties.append(OxmlElement("w:tblHeader"))

    for row_index, values in enumerate(rows):
        for cell_index, value in enumerate(values):
            cell = table.cell(row_index, cell_index)
            shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
            if shading is None:
                shading = OxmlElement("w:shd")
                cell._tc.get_or_add_tcPr().append(shading)
            shading.set(qn("w:fill"), "DCE6F1" if row_index == 0 else "FFFFFF")
            paragraph = cell.paragraphs[0]
            set_paragraph_rtl(paragraph)
            paragraph.paragraph_format.space_after = Pt(0)
            add_rich_text(
                paragraph,
                value,
                size=9.3,
                bold=row_index == 0,
                color=RGBColor(0x1F, 0x4E, 0x79) if row_index == 0 else RGBColor(0x20, 0x2A, 0x35),
            )
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def build_document(source_path: Path, output_path: Path) -> None:
    source_text = source_path.read_text(encoding="utf-8")
    source_hash = hashlib.sha256(source_text.encode("utf-8")).hexdigest()
    blocks = parse_markdown(source_text)
    document = Document()
    configure_styles(document)
    configure_section(document)

    properties = document.core_properties
    properties.title = "ספר פרויקט — Predictor1"
    properties.subject = "RUNI Internet Technologies 2026 final project"
    properties.author = "Predictor1"
    properties.last_modified_by = "Predictor1 deterministic generator"
    properties.comments = f"Generated from docs/project-book-source.md SHA-256 {source_hash}"
    properties.created = FIXED_TIME
    properties.modified = FIXED_TIME
    properties.revision = 1

    current_numbering_id: int | None = None
    for block in blocks:
        if block.kind != "number":
            current_numbering_id = None
        if block.kind == "pagebreak":
            document.add_page_break()
            continue
        if block.kind == "table":
            add_table(document, block.value)
            continue
        if block.kind == "heading1":
            paragraph = document.add_paragraph(style="Title")
        elif block.kind == "heading2":
            paragraph = document.add_paragraph(style="Heading 1")
        elif block.kind == "heading3":
            paragraph = document.add_paragraph(style="Heading 2")
        elif block.kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
        elif block.kind == "number":
            paragraph = document.add_paragraph(style="List Number")
            if current_numbering_id is None:
                current_numbering_id = create_numbering_instance(document)
            apply_numbering(paragraph, current_numbering_id)
        elif block.kind == "callout":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EAF2F8")
            paragraph._p.get_or_add_pPr().append(shading)
        else:
            paragraph = document.add_paragraph()

        set_paragraph_rtl(paragraph)
        if block.kind == "heading1":
            paragraph.paragraph_format.space_before = Pt(18)
        add_rich_text(
            paragraph,
            str(block.value),
            size=11 if block.kind == "callout" else None,
            bold=True if block.kind == "callout" else None,
            color=RGBColor(0x1F, 0x4E, 0x79) if block.kind == "callout" else None,
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="predictor-project-book-") as temp_directory:
        raw_path = Path(temp_directory) / "raw.docx"
        document.save(raw_path)
        with zipfile.ZipFile(raw_path, "r") as source_zip, zipfile.ZipFile(
            output_path,
            "w",
            compression=zipfile.ZIP_STORED,
        ) as output_zip:
            for info in sorted(source_zip.infolist(), key=lambda item: item.filename):
                normalized = zipfile.ZipInfo(info.filename, ZIP_TIME)
                # Stored entries avoid zlib-version-specific DEFLATE bytes while
                # remaining valid OPC/Word packages across supported runtimes.
                normalized.compress_type = zipfile.ZIP_STORED
                normalized.external_attr = info.external_attr
                normalized.create_system = 0
                output_zip.writestr(normalized, source_zip.read(info.filename))


def validate_links(source_path: Path) -> None:
    source_text = source_path.read_text(encoding="utf-8")
    for _label, target in LINK_PATTERN.findall(source_text):
        if target.startswith("https://"):
            continue
        if target.startswith("http://"):
            raise ValueError(f"Project-book link must use HTTPS: {target}")
        resolved = (source_path.parent / target).resolve()
        if not resolved.is_file():
            raise ValueError(f"Project-book local link does not exist: {target}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate the deterministic Predictor1 project book")
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--check", action="store_true")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    source_path = args.source.resolve()
    output_path = args.output.resolve()
    validate_links(source_path)

    if args.check:
        if not output_path.is_file():
            print("Generated project book is missing.", file=sys.stderr)
            return 1
        with tempfile.TemporaryDirectory(prefix="predictor-project-book-check-") as temporary:
            candidate = Path(temporary) / "project-book.docx"
            build_document(source_path, candidate)
            if candidate.read_bytes() != output_path.read_bytes():
                print(
                    "Generated project book is stale. Run: npm run docs:book",
                    file=sys.stderr,
                )
                return 1
        print("Generated project book is current and deterministic.")
        return 0

    build_document(source_path, output_path)
    print(f"Generated {output_path.relative_to(ROOT)} from {source_path.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
